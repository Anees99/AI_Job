"""Resume/CV parser to extract structured data from PDF/DOCX files."""

import logging
import json
import re
from typing import Dict, List, Optional, Any
from pathlib import Path

# Third-party imports
import pdfplumber
from docx import Document

from utils.logger import setup_logger
from utils.config import Config

# Initialize logger
logger = setup_logger(__name__)


class ResumeParser:
    """
    Parse resume/CV files (PDF or DOCX) into structured JSON data.
    
    Extracts:
    - Contact information (name, email, phone, location, LinkedIn)
    - Summary/Objective
    - Skills (categorized if possible)
    - Work experience (company, role, dates, bullet points)
    - Education (degree, institution, dates)
    - Projects
    - Certifications
    """
    
    def __init__(self):
        """Initialize the resume parser."""
        logger.info("ResumeParser initialized")
        self.supported_extensions = ['.pdf', '.docx', '.doc']
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a resume file and return structured data.
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Dictionary containing parsed resume data
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            RuntimeError: If parsing fails
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        extension = file_path.suffix.lower()
        if extension not in self.supported_extensions:
            logger.error(f"Unsupported file format: {extension}")
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.supported_extensions)}"
            )
        
        logger.info(f"Parsing resume: {file_path.name} ({extension})")
        
        try:
            if extension == '.pdf':
                raw_text = self._extract_pdf_text(file_path)
            elif extension in ['.docx', '.doc']:
                raw_text = self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported extension: {extension}")
            
            if not raw_text or len(raw_text.strip()) < 50:
                logger.warning("Extracted text is too short")
                raise RuntimeError("Could not extract sufficient text from resume")
            
            # Parse structured data from raw text
            structured_data = self._parse_content(raw_text)
            
            # Add metadata
            structured_data['metadata'] = {
                'source_file': str(file_path),
                'file_name': file_path.name,
                'character_count': len(raw_text),
                'word_count': len(raw_text.split())
            }
            
            logger.info(
                f"Successfully parsed resume: {structured_data.get('contact', {}).get('name', 'Unknown')} "
                f"with {len(structured_data.get('experience', []))} work experiences"
            )
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Failed to parse resume: {str(e)}")
            raise RuntimeError(f"Failed to parse resume: {str(e)}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        Extract text from a PDF file using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        logger.debug(f"Extracting text from PDF: {file_path}")
        
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        logger.debug(f"Page {page_num}: extracted {len(page_text)} chars")
            
            full_text = '\n\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        logger.debug(f"Extracting text from DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            
            full_text = '\n\n'.join(paragraphs)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise
    
    def _parse_content(self, text: str) -> Dict[str, Any]:
        """
        Parse raw text into structured resume data.
        
        Uses pattern matching and heuristics to identify sections.
        
        Args:
            text: Raw resume text
            
        Returns:
            Structured dictionary with resume sections
        """
        logger.debug("Parsing resume content into structured data")
        
        lines = text.split('\n')
        lines = [line.strip() for line in lines if line.strip()]
        
        data = {
            'contact': self._extract_contact_info(text),
            'summary': self._extract_summary(text, lines),
            'skills': self._extract_skills(text),
            'experience': self._extract_experience(text, lines),
            'education': self._extract_education(text, lines),
            'projects': self._extract_projects(text, lines),
            'certifications': self._extract_certifications(text, lines),
            'languages': self._extract_languages(text)
        }
        
        return data
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume text."""
        contact = {
            'name': None,
            'email': None,
            'phone': None,
            'location': None,
            'linkedin': None,
            'website': None
        }
        
        # Email pattern
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group(0)
        
        # Phone pattern (various formats)
        phone_patterns = [
            r'\+?[\d\s\-\(\)]{10,}',
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
            r'\d{3}[-.]?\d{3}[-.]?\d{4}'
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact['phone'] = phone_match.group(0).strip()
                break
        
        # LinkedIn URL
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group(0)
        
        # Website/Portfolio
        website_pattern = r'(?:https?://)?(?:www\.)?[a-zA-Z0-9\-]+\.[a-z]{2,}(?:/[^\s]*)?'
        website_matches = re.findall(website_pattern, text, re.IGNORECASE)
        for url in website_matches:
            if 'linkedin' not in url.lower() and '@' not in url:
                contact['website'] = url
                break
        
        # Name (usually at the top, capitalized words)
        first_lines = text.split('\n')[:5]
        for line in first_lines:
            line = line.strip()
            # Skip lines with email, phone, or URLs
            if re.search(r'@|http|\d{3}', line):
                continue
            # Assume first substantial line without special chars is name
            if len(line) > 2 and len(line) < 50 and re.match(r'^[A-Za-z\s\.]+$', line):
                contact['name'] = line
                break
        
        # Location (city, state pattern)
        location_pattern = r'[A-Z][a-z]+,\s*[A-Z]{2}(?:\s+\d{5})?'
        location_match = re.search(location_pattern, text)
        if location_match:
            contact['location'] = location_match.group(0)
        
        logger.debug(f"Extracted contact: {contact['name']}, {contact['email']}")
        return contact
    
    def _extract_summary(self, text: str, lines: List[str]) -> Optional[str]:
        """Extract professional summary/objective."""
        summary_keywords = [
            'summary', 'objective', 'profile', 'about', 
            'professional summary', 'career objective'
        ]
        
        summary_start = -1
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in summary_keywords):
                summary_start = i + 1
                break
        
        if summary_start == -1:
            # Try to find paragraph after name/contact
            for i, line in enumerate(lines[3:8], 3):
                if len(line) > 50 and len(line) < 500:
                    summary_start = i
                    break
        
        if summary_start != -1:
            # Collect summary lines until next section
            summary_lines = []
            for i in range(summary_start, min(summary_start + 10, len(lines))):
                line = lines[i]
                # Stop if we hit a new section header
                if line.isupper() or line.endswith(':') or len(line) < 20:
                    break
                summary_lines.append(line)
            
            if summary_lines:
                return ' '.join(summary_lines)
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume."""
        skills = []
        
        # Look for skills section
        skills_section_patterns = [
            r'(?:SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|EXPERTISE)[:\s]*(.*?)(?=\n[A-Z]{2,}|$)',
            r'(?:Skills:?)\s*(.*?)(?=\n[A-Z]|$)'
        ]
        
        for pattern in skills_section_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                skills_text = match.group(1)
                # Split by common delimiters
                skill_list = re.split(r'[,;|•\-\n]', skills_text)
                skills.extend([s.strip() for s in skill_list if s.strip() and len(s.strip()) < 50])
                break
        
        # If no dedicated section, look for skill-like patterns
        if not skills:
            # Common tech skills pattern
            tech_patterns = [
                r'\b(Python|Java|JavaScript|C\+\+|SQL|AWS|Azure|Docker|Kubernetes|React|Node\.js|TensorFlow|PyTorch)\b',
                r'\b(Agile|Scrum|Machine Learning|Data Analysis|Project Management)\b'
            ]
            for pattern in tech_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                skills.extend(matches)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower not in seen and len(skill) > 1:
                seen.add(skill_lower)
                unique_skills.append(skill)
        
        return unique_skills[:30]  # Limit to 30 skills
    
    def _extract_experience(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract work experience entries."""
        experiences = []
        
        # Keywords that indicate experience section
        exp_keywords = ['experience', 'employment', 'work history', 'professional experience']
        
        exp_start = -1
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in exp_keywords):
                exp_start = i
                break
        
        if exp_start == -1:
            exp_start = 0
        
        # Simple heuristic: look for company/role patterns
        current_exp = {}
        date_pattern = r'(\d{1,2}/\d{2,4}|\w+ \d{4}|Present|Current|-\s*Present)'
        
        i = exp_start
        while i < len(lines):
            line = lines[i].strip()
            
            # Check if this looks like a job title/company line
            if len(line) > 3 and len(line) < 80:
                # Check for date pattern nearby
                next_few_lines = ' '.join(lines[i:i+3])
                if re.search(date_pattern, next_few_lines):
                    if current_exp:
                        experiences.append(current_exp)
                    
                    current_exp = {
                        'title': line,
                        'company': '',
                        'dates': '',
                        'bullets': []
                    }
                    
                    # Try to extract company from next line
                    if i + 1 < len(lines) and len(lines[i+1]) < 60:
                        current_exp['company'] = lines[i+1]
                        i += 1
                    
                    # Extract dates
                    date_match = re.search(date_pattern, next_few_lines)
                    if date_match:
                        current_exp['dates'] = date_match.group(0)
            
            # Collect bullet points
            elif current_exp and (line.startswith('•') or line.startswith('-') or 
                                  line.startswith('●') or (len(line) > 20 and len(line) < 200)):
                if len(line) > 15:  # Meaningful bullet point
                    current_exp['bullets'].append(line.lstrip('•-● '))
            
            i += 1
        
        if current_exp:
            experiences.append(current_exp)
        
        logger.debug(f"Extracted {len(experiences)} work experiences")
        return experiences
    
    def _extract_education(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract education entries."""
        education = []
        
        edu_keywords = ['education', 'academic background', 'qualifications', 'degree']
        
        edu_start = -1
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in edu_keywords):
                edu_start = i
                break
        
        if edu_start == -1:
            return education
        
        # Look for degree patterns
        degree_patterns = [
            r'(Bachelor[\'s]?|Master[\'s]?|Ph\.?D\.?|MBA|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?)',
            r'(Bachelors|Masters|Doctorate)'
        ]
        
        current_edu = {}
        for i in range(edu_start, len(lines)):
            line = lines[i].strip()
            
            for pattern in degree_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    if current_edu:
                        education.append(current_edu)
                    current_edu = {
                        'degree': line,
                        'institution': '',
                        'dates': '',
                        'details': []
                    }
                    
                    # Try to get institution from next line
                    if i + 1 < len(lines) and len(lines[i+1]) < 80:
                        current_edu['institution'] = lines[i+1]
                    
                    # Get dates
                    date_match = re.search(r'(\d{4}(?:\s*[-–]\s*\d{4}|[-–]Present)?)', line)
                    if date_match:
                        current_edu['dates'] = date_match.group(1)
                    break
        
        if current_edu:
            education.append(current_edu)
        
        return education
    
    def _extract_projects(self, text: str, lines: List[str]) -> List[Dict[str, Any]]:
        """Extract project entries."""
        projects = []
        
        proj_keywords = ['projects', 'personal projects', 'key projects']
        
        proj_start = -1
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in proj_keywords):
                proj_start = i
                break
        
        if proj_start == -1:
            return projects
        
        current_proj = {}
        for i in range(proj_start + 1, len(lines)):
            line = lines[i].strip()
            
            if len(line) > 5 and len(line) < 100 and not line.startswith('•'):
                if current_proj:
                    projects.append(current_proj)
                current_proj = {
                    'name': line,
                    'description': [],
                    'technologies': []
                }
            elif current_proj and line.startswith('•'):
                desc = line.lstrip('• ')
                # Check for technologies mentioned
                tech_match = re.search(
                    r'(Python|Java|JavaScript|React|Node|SQL|AWS|Docker|etc\.)',
                    desc, re.IGNORECASE
                )
                if tech_match:
                    current_proj['technologies'].append(tech_match.group(1))
                current_proj['description'].append(desc)
        
        if current_proj:
            projects.append(current_proj)
        
        return projects
    
    def _extract_certifications(self, text: str, lines: List[str]) -> List[str]:
        """Extract certifications."""
        certs = []
        
        cert_keywords = ['certification', 'license', 'credential']
        
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in cert_keywords):
                # Clean up the line
                cert = re.sub(r'^(Certifications?|Licenses?):?\s*', '', line, flags=re.IGNORECASE)
                if cert and len(cert) > 5:
                    certs.append(cert.strip('•-: '))
        
        return certs
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extract languages spoken."""
        languages = []
        
        lang_keywords = ['languages', 'language proficiency']
        
        # Simple extraction of common languages
        common_languages = [
            'English', 'Spanish', 'French', 'German', 'Chinese', 'Japanese',
            'Korean', 'Portuguese', 'Italian', 'Russian', 'Arabic', 'Hindi'
        ]
        
        for lang in common_languages:
            if re.search(rf'\b{lang}\b', text, re.IGNORECASE):
                languages.append(lang)
        
        return languages
    
    def save_to_json(self, data: Dict[str, Any], output_path: str) -> str:
        """
        Save parsed resume data to JSON file.
        
        Args:
            data: Parsed resume dictionary
            output_path: Path to save JSON file
            
        Returns:
            Path to saved file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved parsed resume to {output_path}")
        return str(output_path)
    
    def load_from_json(self, json_path: str) -> Dict[str, Any]:
        """
        Load parsed resume data from JSON file.
        
        Args:
            json_path: Path to JSON file
            
        Returns:
            Parsed resume dictionary
        """
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        logger.info(f"Loaded parsed resume from {json_path}")
        return data
